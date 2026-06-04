"""
generate_project_report.py
--------------------------
Generates  outputs/FinnUp_Project_Report.docx
  – APAL Cohort 2 | Group 1 | IIM Calcutta
    Full Project Status, Post-EDA Action Plan & Timeline to May 2026

Run:
    python generate_project_report.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = Path("outputs") / "FinnUp_Project_Report.docx"
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

# ── Colour helpers ─────────────────────────────────────────────────────────────
def hex_to_rgb(h): return tuple(int(h[i:i+2],16) for i in (0,2,4))

NAVY   = "1B3A6B"
TEAL   = "0D9488"
GREEN  = "16A34A"
AMBER  = "F59E0B"
RED    = "EF4444"
PURPLE = "7C3AED"
GRAY   = "64748B"
LTBLUE = "DBEAFE"
WHITE  = "FFFFFF"
LGRAY  = "F5F7FA"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get(side, 'none'))
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def para_color(para, hex_color):
    for run in para.runs:
        run.font.color.rgb = RGBColor(*hex_to_rgb(hex_color))

def add_cell_para(cell, text, bold=False, size=10, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    return p

# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Normal style
for style_name in ['Normal']:
    style = doc.styles[style_name]
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

# ── Helper functions ────────────────────────────────────────────────────────────
def heading(text, level=1, color=NAVY, space_before=12, space_after=4):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))
        run.font.name = 'Calibri'
    return p

def para(text='', bold=False, italic=False, color=None, size=10.5,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    return p

def bullet(text, level=0, color=None, bold_prefix=None, prefix_color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix + ' ')
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        if prefix_color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(prefix_color))
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.name = 'Calibri'
    if color:
        run2.font.color.rgb = RGBColor(*hex_to_rgb(color))
    return p

def add_banner(text, bg=NAVY, fg=WHITE, size=13):
    """Shaded paragraph acting as a section banner."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.3)
    set_cell_bg(cell, bg)
    p = add_cell_para(cell, text, bold=True, size=size, color=fg)
    cell._tc.get_or_add_tcPr()
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()
    return tbl

def add_info_box(title, title_bg, items):
    """Shaded table: header row + bullet rows."""
    tbl = doc.add_table(rows=1 + len(items), cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    hdr = tbl.rows[0].cells[0]
    set_cell_bg(hdr, title_bg)
    add_cell_para(hdr, title, bold=True, size=11, color=WHITE)
    # Items
    for i, item in enumerate(items):
        c = tbl.rows[i+1].cells[0]
        set_cell_bg(c, LGRAY if i % 2 == 0 else WHITE)
        add_cell_para(c, f"  ▸  {item}", size=10, color="1E293B")
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), TEAL)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(20)
run = p_title.add_run("FinnUp MSME Lender Matching")
run.bold = True; run.font.size = Pt(22); run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Project Status Report, Post-EDA Action Plan & Delivery Roadmap")
run2.font.size = Pt(14); run2.font.name = 'Calibri'
run2.font.color.rgb = RGBColor(*hex_to_rgb(TEAL))

para()
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_meta.add_run(
    "APAL Programme  |  IIM Calcutta  |  Cohort 2  |  Group 1\n"
    "April 2026  |  Target completion: 31 May 2026"
)
run3.font.size = Pt(11); run3.font.name = 'Calibri'
run3.font.color.rgb = RGBColor(*hex_to_rgb(GRAY))

para()
p_mem = doc.add_paragraph()
p_mem.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p_mem.add_run(
    "Asha · Arvind · Anil · Bhupesh · Deepak · Ganesh · Gopal\n"
    "Hareram · Pranali · Rahul · Savitha · Sonam · Samik"
)
run4.font.size = Pt(10); run4.font.name = 'Calibri'
run4.font.color.rgb = RGBColor(*hex_to_rgb(GRAY))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – WHAT IS COMPLETED
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 1 — WHAT IS COMPLETED  ✓", bg=NAVY)

heading("1.1  Phase 0 — Problem Definition & Proposal (January 2026)", level=2)
para("Submitted to IIM Calcutta on 18 January 2026. The following was delivered:")
completed_p0 = [
    "Identified and framed the core business problem: FinnUp's rule-based engine achieves only 8–10% lead-to-conversion; 80–85% of sales effort is wasted on mismatched borrower-lender pairs",
    "Documented the four root causes: lack of intelligent filtering, late policy mismatch discovery, high rejection rates, and poor borrower experience",
    "Proposed a two-stage ML solution: Stage 1 predicts lender-specific approval probability (XGBoost); Stage 2 ranks eligible lenders by predicted acceptance",
    "Designed the full AI Canvas: Prediction → Judgement → Action → Outcome",
    "Defined the Continuous Learning Loop: every approval/rejection outcome retrains the model",
    "Set business targets: 11% → 75% conversion rate, 5× revenue by Year 3, 20% ROI in 2 years",
    "Quantified the market opportunity: 50 Lakh+ MSME leads, ₹20 lakh average ticket size, 10,000 qualified leads annually on FinnUp",
]
for item in completed_p0:
    bullet(item, bold_prefix="✓", prefix_color=GREEN)

divider()

heading("1.2  Phase 1 — Data Acquisition & Exploratory Data Analysis (Feb–Apr 2026)", level=2)
para("Data was extracted directly from FinnUp's production database and exported to Excel. A comprehensive EDA was completed across all 11 data sheets.")

heading("1.2.1  Data Secured", level=3, color=TEAL, space_before=6)
completed_data = [
    "Real production data extracted from FinnUp's live database — not sample or synthetic data",
    "11 structured data sheets covering the full MSME borrower lifecycle",
    "6,483 unique MSME borrower profiles loaded and validated",
    "9,273 director records (average 1.4 directors per borrower)",
    "3,996 bank statement rows covering 2,489 borrowers (38% of total)",
    "Multi-year financial data: Balance Sheet, P&L, Cash Flow, Financial KPIs, Financial Summary",
    "Loan request history: product applications, amounts, FinnUp status per borrower",
    "Document completeness and reference data captured",
]
for item in completed_data:
    bullet(item, bold_prefix="✓", prefix_color=GREEN)

heading("1.2.2  EDA Completed", level=3, color=TEAL, space_before=6)
completed_eda = [
    "Missing value analysis across all 50+ columns of the Borrower Profile sheet — identified 16 columns with >50% missing",
    "CIBIL score analysis: only 55 borrowers (~0.8%) have actual borrower CIBIL scores — identified as critical data gap",
    "Director CIBIL profiled: 86% coverage across 9,273 director records — confirmed as primary credit proxy",
    "Turnover distribution: 100% coverage, log-normally distributed (median log ≈ 16.8 ≈ ₹20M range) — most reliable numeric feature",
    "Networth distribution: 92% coverage, right-skewed — reliable feature after log transformation",
    "Business age distribution: 51% coverage — usable with imputation",
    "Company type distribution: Proprietorships dominate at >60%, followed by Private Limited companies",
    "Geographic distribution: state data available for 55% of borrowers; top 5 states identified",
    "Bank risk signals quantified: Bounce Ratio, Credit-Debit Ratio, Avg EOD Balance, Inward/Outward Cheque Bounces",
    "Label landscape mapped: FinnUp Status (best available binary label), Lead Status (99%+ unlabelled), Count Disbursed > 0 (6% of borrowers)",
    "Financial KPI analysis: DSCR, Current Ratio, Net Profit Margin, Debt-Equity Ratio — available for only 500 borrowers (7.7%)",
    "Correlation matrix computed across 30+ numeric features — key co-linear groups identified",
    "30+ features identified as production-ready for ML model training",
    "Full EDA automated as a self-regenerating Python notebook and PowerPoint report",
]
for item in completed_eda:
    bullet(item, bold_prefix="✓", prefix_color=GREEN)

divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 – KEY EDA FINDINGS & WHAT THEY MEAN
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 2 — KEY EDA FINDINGS & IMPLICATIONS", bg=TEAL)

heading("2.1  Critical Data Gaps Found", level=2)

para("The EDA revealed several important gaps that must be addressed before model training:", space_after=4)

gaps = [
    ("Borrower CIBIL Score: Only 0.8% coverage (55 of 6,483 borrowers)",
     "Action: Use Director CIBIL (86% coverage) as proxy; build a CIBIL imputation model"),
    ("Bank Statement Data: Only 38% of borrowers have bank data (2,489 of 6,483)",
     "Action: Expand bank data collection via Account Aggregator (AA) consent; use as optional feature with missingness flag"),
    ("Financial KPI Data: Only 7.7% coverage (500 borrowers)",
     "Action: Use as a supplementary feature with 'KPI available' binary flag; impute with sector-level medians"),
    ("Lead Status Label: 99%+ of borrowers are unlabelled",
     "Action: Use FinnUp Status (from Loan Requests sheet) as primary training label"),
    ("State / Location: 45% missing",
     "Action: Impute from pin code or city if available; use 'state_unknown' as a category"),
    ("Business Age: 51% missing",
     "Action: Impute using registration date if available; else use median within company type group"),
    ("Director Ownership %: Partially missing",
     "Action: Impute with 100% for sole proprietors; use 0 for missing with missingness flag"),
]

tbl_gaps = doc.add_table(rows=1 + len(gaps), cols=2)
tbl_gaps.style = 'Table Grid'
tbl_gaps.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr_cells = tbl_gaps.rows[0].cells
for c, txt in zip(hdr_cells, ["Data Gap Identified", "Recommended Action"]):
    set_cell_bg(c, NAVY)
    add_cell_para(c, txt, bold=True, size=10.5, color=WHITE)

for i, (gap, action) in enumerate(gaps):
    row = tbl_gaps.rows[i + 1]
    set_cell_bg(row.cells[0], LGRAY if i % 2 == 0 else WHITE)
    set_cell_bg(row.cells[1], LGRAY if i % 2 == 0 else WHITE)
    add_cell_para(row.cells[0], gap, size=10, color="1E293B")
    add_cell_para(row.cells[1], action, size=10, color="1E293B")

para()

heading("2.2  The Small Dataset Problem — 6,483 Records Is Manageable But Needs Strategy", level=2)

para(
    "6,483 borrowers is a relatively small dataset for training a robust ML model. "
    "However, this is not unusual in MSME lending, and there are well-established strategies to address it.",
    space_after=6
)

para("Why 6,483 records is a concern:", bold=True, color=RED, space_after=2)
small_data_problems = [
    "Classification models (XGBoost) can overfit with fewer than 5,000 records if features are many",
    "Class imbalance is likely: very few 'approved' vs 'applied' records — need SMOTE or class weighting",
    "Held-out test set will be small (~1,300 records), making AUC confidence intervals wide",
    "Lender-specific models (per lender) would have even fewer training examples",
]
for item in small_data_problems:
    bullet(item, bold_prefix="⚠", prefix_color=AMBER)

para()
para("Strategies to address the small dataset:", bold=True, color=TEAL, space_after=2)

strategies = [
    ("Expand FinnUp's data coverage",
     "Request additional borrower records from FinnUp — any new registrations since the extract date should be included. Target: 10,000+ records for reliable model training."),
    ("Use the full multi-year transaction history",
     "Each borrower may have multiple loan applications across years. Treat each application as an independent training row — this can multiply usable rows to 15,000–20,000 records."),
    ("Include rejection data from lender partners",
     "Request approval/rejection logs from lender partners — even anonymised outcomes add critical labelled rows."),
    ("Transfer learning from similar MSME datasets",
     "Fine-tune a model pre-trained on public MSME credit datasets (e.g., SIDBI, RBI data), then fine-tune on FinnUp data."),
    ("Synthetic data generation (SMOTE / CTGAN)",
     "For the minority class (approved loans), use SMOTE to generate synthetic approved-borrower samples during training. Use with caution — validate on real data only."),
    ("Feature-rich shallow models work well on small data",
     "XGBoost with L1/L2 regularisation and careful feature selection performs reliably on 5,000–10,000 records. Avoid deep learning — it requires 100,000+ records."),
    ("K-fold cross-validation",
     "Use 5-fold stratified cross-validation instead of a single train-test split to maximise learning from all available records."),
    ("Lender-agnostic model initially",
     "Train one global model (not per lender) to maximise training data. Apply lender policy rules as a post-scoring filter layer."),
]

tbl_strat = doc.add_table(rows=1 + len(strategies), cols=2)
tbl_strat.style = 'Table Grid'
hdr_s = tbl_strat.rows[0].cells
for c, txt in zip(hdr_s, ["Strategy", "Implementation Detail"]):
    set_cell_bg(c, TEAL)
    add_cell_para(c, txt, bold=True, size=10.5, color=WHITE)
for i, (strat, detail) in enumerate(strategies):
    row = tbl_strat.rows[i + 1]
    set_cell_bg(row.cells[0], LGRAY if i % 2 == 0 else WHITE)
    set_cell_bg(row.cells[1], LGRAY if i % 2 == 0 else WHITE)
    add_cell_para(row.cells[0], strat, bold=True, size=10, color=NAVY)
    add_cell_para(row.cells[1], detail, size=10, color="1E293B")

para()
divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 – POST-EDA NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 3 — POST-EDA NEXT STEPS (DETAILED ACTION PLAN)", bg=PURPLE)

heading("3.1  Immediate Actions (Week 1 — This Week)", level=2)

para("These must be done before any modelling can proceed:", bold=True, color=RED, space_after=4)

immediate = [
    ("BLOCKER — Get outcome labels from FinnUp",
     "Request the CRM/LOS export of formal approved/rejected/disbursed status per borrower–loan pair. "
     "Without this, supervised model training cannot begin. Even 1,000 labelled records are enough to start. "
     "FinnUp Status (YES/NO) in the Loan Requests sheet will be used as proxy in the meantime."),
    ("BLOCKER — Confirm lender policy files are complete",
     "Verify that the policy JSON/Excel files in /data/lender_policies/ are current and cover all active lenders. "
     "Each lender's minimum CIBIL, turnover threshold, sector restrictions, and ticket size range is needed."),
    ("Get more borrower records if available",
     "Request any new borrower registrations since the original export. "
     "Also request historical data going back 3+ years if available — more rows = more training data."),
    ("Clarify bank statement coverage",
     "Confirm whether Account Aggregator (AA) consent data is available for the remaining 62% of "
     "borrowers who do not have bank statement data. If yes, add to pipeline."),
]
for title, detail in immediate:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"[CRITICAL]  {title}: ")
    run_t.bold = True; run_t.font.size = Pt(10.5); run_t.font.name = 'Calibri'
    run_t.font.color.rgb = RGBColor(*hex_to_rgb(RED))
    run_d = p.add_run(detail)
    run_d.font.size = Pt(10.5); run_d.font.name = 'Calibri'

para()

heading("3.2  Feature Engineering Plan (Weeks 1–2)", level=2)
para("Transform raw data from all 11 sheets into 50+ model-ready features:", space_after=4)

feature_groups = [
    ("A. Profile Features (from Borrower Profile sheet)", [
        "log(Turnover) — log-transform to handle right skew; treat 0 separately",
        "log(Networth) — same transformation; flag negatives separately as 'negative_networth'",
        "Business Age (years) — bin into: <2yr, 2–5yr, 5–10yr, 10yr+; impute missing with group median",
        "Company Type — one-hot encode (Proprietorship, Partnership, Pvt Ltd, LLP, Public Ltd, Other)",
        "Industry / Sector — group into 6 macro-sectors; label encode",
        "State — encode top-20 states; group rare states as 'Other'",
        "Documents completeness score — count key documents present / total expected",
        "Application-to-first-response gap (days) — if date fields available",
    ]),
    ("B. Director / Credit Features (from Directors sheet)", [
        "min_director_cibil — minimum CIBIL score across all directors (most conservative signal)",
        "max_director_cibil — maximum CIBIL score across all directors",
        "avg_director_cibil — average CIBIL across directors",
        "director_cibil_below_650 — binary flag: any director with CIBIL < 650",
        "num_directors — count of directors (more directors = more complex governance)",
        "max_ownership_pct — highest individual ownership percentage (promoter concentration)",
        "cibil_imputed_flag — binary: 1 if borrower CIBIL was imputed from director proxy",
        "imputed_borrower_cibil — regression imputation using director CIBIL + financial features",
    ]),
    ("C. Bank Statement Features (from Bank Statements sheet)", [
        "bounce_ratio — total cheque bounces / total credits (key risk signal)",
        "bounce_ratio_missing — binary flag: 1 if no bank data available",
        "log_avg_eod_balance — log-transformed average end-of-day balance",
        "credit_debit_ratio — total credits / total debits (>1 = net inflow)",
        "inward_cheque_bounces — count of inward bounces (capped at 30)",
        "outward_cheque_bounces — count of outward bounces (capped at 30)",
        "balance_volatility — std deviation of monthly EOD balance / mean (stability indicator)",
        "bank_vintage_months — number of months of bank history available",
        "has_bank_data — binary flag: 1 if bank statement data exists for borrower",
    ]),
    ("D. Loan Request Features (from Loan Requests sheet)", [
        "loan_amount_requested — amount requested (log-transformed)",
        "num_loan_applications — total number of applications submitted",
        "num_products_tried — distinct loan products applied for",
        "days_since_first_application — borrower tenure on platform",
        "finnup_status_any_yes — binary: 1 if any loan application has FinnUp Status = YES",
        "product_category — macro-category of primary loan product (working capital, term loan, etc.)",
    ]),
    ("E. Financial KPI Features (from Financial KPIs sheet — 500 borrowers)", [
        "dscr_avg — Debt Service Coverage Ratio (higher = better repayment capacity)",
        "current_ratio — current assets / current liabilities (liquidity signal)",
        "debt_equity_ratio — leverage measure",
        "net_profit_margin_pct — profitability indicator",
        "pbdita_margin_pct — EBITDA proxy margin",
        "revenue_yoy_growth — year-on-year revenue growth rate",
        "kpi_data_available — binary flag: 1 if formal financial KPIs exist",
        "All KPI features imputed with sector median for borrowers without financials",
    ]),
    ("F. Derived / Engineered Features", [
        "Borrower-to-director CIBIL gap — abs(borrower_cibil - min_director_cibil)",
        "Size tier — micro / small / medium based on turnover bands",
        "Credit risk score (rule-based) — simple weighted score from CIBIL + bounce + CDR",
        "Data quality score — fraction of key fields that are non-null for that borrower",
        "Lender eligibility count — number of lenders whose hard rules this borrower passes",
    ]),
]

for group_title, features in feature_groups:
    heading(group_title, level=3, color=TEAL, space_before=6, space_after=2)
    for f in features:
        bullet(f, level=1)

para()
divider()
doc.add_page_break()

heading("3.3  Missing Value Treatment Plan", level=2)
para("Handling missing data correctly is critical with only 6,483 records — we cannot afford to drop rows.", space_after=4)

mv_table = [
    ("CIBIL Score (Borrower)", "0.8%", "Regression imputation using Director CIBIL + Turnover + Business Age. Flag as 'imputed'"),
    ("Business Age", "51%", "Impute with median within same Company Type group. Add 'business_age_imputed' binary flag"),
    ("State / Location", "45%", "Assign 'Unknown' category. Use as a feature with 'state_unknown' level"),
    ("Director CIBIL", "14%", "Impute with median CIBIL by company type + sector. Flag as 'director_cibil_imputed'"),
    ("Bank Data (all bank features)", "62%", "Set all bank features to 0. Add 'has_bank_data' = 0 flag. This is meaningful — absence itself is a risk signal"),
    ("Director Ownership %", "~20%", "For sole proprietors: impute as 100%. For others: impute with 51% (majority threshold). Add flag"),
    ("Financial KPIs (DSCR, ratios)", "92%", "Impute with sector-level median. Add 'kpi_available' = 0 flag"),
    ("Networth", "8%", "Impute with sector + size tier median. Do not drop rows"),
    ("Loan Amount Requested", "Partial", "Impute with product-category median amount. Flag as missing"),
]

tbl_mv = doc.add_table(rows=1 + len(mv_table), cols=3)
tbl_mv.style = 'Table Grid'
for c, txt in zip(tbl_mv.rows[0].cells, ["Feature / Column", "Missing %", "Treatment Strategy"]):
    set_cell_bg(c, NAVY)
    add_cell_para(c, txt, bold=True, size=10.5, color=WHITE)
for i, (feat, pct, treatment) in enumerate(mv_table):
    row = tbl_mv.rows[i + 1]
    set_cell_bg(row.cells[0], LGRAY if i % 2 == 0 else WHITE)
    set_cell_bg(row.cells[1], LGRAY if i % 2 == 0 else WHITE)
    set_cell_bg(row.cells[2], LGRAY if i % 2 == 0 else WHITE)
    add_cell_para(row.cells[0], feat, bold=True, size=10, color=NAVY)
    add_cell_para(row.cells[1], pct, size=10, color="1E293B", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_para(row.cells[2], treatment, size=10, color="1E293B")

para()
para("Key principle: Every missing value is replaced by an imputed value PLUS a binary 'field_missing' flag feature. "
     "This allows the model to learn both the imputed value AND the fact that data was missing.",
     bold=True, color=TEAL, space_after=6)

divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 – FULL TIMELINE TO MAY 2026
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 4 — FULL DELIVERY TIMELINE  (April – May 2026)", bg=NAVY)

heading("Overview", level=2)
para(
    "Project started January 2026. EDA is complete as of early April 2026. "
    "All remaining phases must be completed by 31 May 2026 — 8 weeks from today (4 April 2026). "
    "The timeline below is structured week-by-week with clear owners and deliverables.",
    space_after=6
)

weeks = [
    # (week_label, dates, phase, tasks, owner, deliverable, risk)
    (
        "Week 1",
        "7–11 Apr",
        "Data Completion\n& Label Acquisition",
        [
            "Collect formal outcome labels from FinnUp CRM/LOS",
            "Request any additional borrower records (target: 10,000+)",
            "Confirm lender policy files are complete and current",
            "Clarify AA/bank consent status for 62% uncovered borrowers",
            "Review multi-year loan application history — expand training rows",
        ],
        "Ganesh / FinnUp Team",
        "Outcome label dataset (CSV); augmented borrower records; confirmed lender policies",
        "FinnUp team responsiveness — if labels not received, use FinnUp Status proxy and proceed"
    ),
    (
        "Week 2",
        "14–18 Apr",
        "Feature Engineering\n(Phase 2)",
        [
            "Build notebook 02_feature_engineering.ipynb",
            "Engineer all Profile features (log transforms, binning, encoding)",
            "Engineer Director CIBIL features (min/max/avg, flags)",
            "Build CIBIL imputation model (regression on director + financial features)",
            "Engineer Bank Statement features (bounce ratio, CDR, EOD balance)",
            "Implement all missing value imputation strategies",
        ],
        "Pranali / Rahul / Anil",
        "Feature engineering notebook; feature matrix with 50+ columns; imputation pipeline",
        "Data quality issues or unexpected column name mismatches in raw sheets"
    ),
    (
        "Week 3",
        "21–25 Apr",
        "Feature Engineering\n& Label Creation\n(Phase 2 cont.)",
        [
            "Engineer Loan Request features (product type, amount, history)",
            "Engineer Financial KPI features with sector-median imputation",
            "Build all derived / cross-sheet features",
            "Create final training labels (binary: approved=1 / not approved=0)",
            "Train-test split: temporal holdout (last 3 months as test set)",
            "Produce feature importance analysis (correlation, mutual information)",
            "Feature selection: remove near-zero variance and highly correlated features",
        ],
        "Pranali / Rahul / Deepak",
        "Final feature matrix; label series; train/test split files; feature importance report",
        "Label imbalance — if <10% positive class, SMOTE must be applied"
    ),
    (
        "Week 4",
        "28 Apr – 2 May",
        "Baseline Model\n& XGBoost Stage 1\n(Phase 3)",
        [
            "Train Logistic Regression baseline on top-10 features — establish AUC-ROC benchmark",
            "Train XGBoost Stage 1 classifier (predict P(approval) per borrower-lender pair)",
            "Apply class imbalance handling: SMOTE oversampling + scale_pos_weight",
            "Hyperparameter tuning via Optuna (5-fold stratified CV)",
            "Feature importance analysis via SHAP values",
            "Validate: AUC-ROC, Precision, Recall, F1 on holdout test set",
            "Fairness check: score distributions by company type, state, sector",
        ],
        "Rahul / Ganesh",
        "Trained baseline + XGBoost model; AUC-ROC report; SHAP feature importance chart; model comparison table",
        "AUC < 0.70 would indicate label quality issue — investigate FinnUp Status label reliability"
    ),
    (
        "Week 5",
        "5–9 May",
        "Stage 2 Lender Ranking\n& Model Card\n(Phase 3 cont.)",
        [
            "Build Stage 2: weighted lender ranking model using Stage 1 output + lender constraints",
            "Implement lender policy rule engine (hard eligibility filter before ML)",
            "Generate SHAP-based explanations per prediction ('Top 5 reasons')",
            "Produce credit score 0–1000 from calibrated probability (Platt scaling)",
            "Define risk bands: Poor (<400), Fair (400–600), Good (600–750), Very Good (750–900), Excellent (900+)",
            "Write model card: performance metrics, fairness analysis, known limitations",
            "Run end-to-end test on 100 sample borrowers",
        ],
        "Rahul / Gopal / Samik",
        "Stage 2 ranking model; lender rule engine; score band definitions; model card document",
        "Lender policy data gaps — fall back to single global threshold if per-lender insufficient"
    ),
    (
        "Week 6",
        "12–16 May",
        "API Development\n(Phase 4)",
        [
            "Build FastAPI REST endpoint: POST /v1/score (single borrower scoring)",
            "Build POST /v1/score/batch (bulk CSV/JSON scoring)",
            "Add GET /v1/lenders/match (top lenders for a borrower)",
            "Implement input validation (Pydantic schemas), API key authentication",
            "Containerise with Docker (Dockerfile + docker-compose)",
            "Write API documentation (auto-generated OpenAPI / Swagger)",
            "Internal testing: verify API returns correct scorecard JSON",
        ],
        "Ganesh / Samik",
        "Working FastAPI application in Docker; API documentation; test results",
        "Environment/dependency issues — use virtual environment and pin all versions"
    ),
    (
        "Week 7",
        "19–23 May",
        "FinnUp Integration\n& UAT (Phase 4 cont.)",
        [
            "Deploy API to cloud (AWS/Azure — containerised)",
            "Configure FinnUp integration: webhook or batch scoring call",
            "Run UAT with FinnUp team: score 200 real borrowers, validate output format",
            "Scorecard format sign-off: score, risk band, top 5 drivers, eligible lenders list",
            "A/B comparison setup: ML-scored leads vs current rule-based leads",
            "Fix any issues found during UAT feedback",
        ],
        "Ganesh / Samik / FinnUp Tech",
        "Live API endpoint; UAT results report; FinnUp integration working",
        "FinnUp tech team availability for integration — schedule coordination needed immediately"
    ),
    (
        "Week 8",
        "26–30 May",
        "Final Report &\nIIM Submission\n(Phase 5)",
        [
            "Set up monitoring: score drift alerts, data quality dashboard, API health check",
            "Document model refresh cycle (monthly retraining plan)",
            "Write final project report for IIM Calcutta (this document, updated with results)",
            "Prepare final PowerPoint presentation (updated EDA + model results)",
            "Conduct stakeholder demo with FinnUp team",
            "Submit all deliverables to IIM Calcutta APAL programme",
            "Capture learning reflections for academic submission",
        ],
        "All team members",
        "Final IIM Calcutta submission package: report, slides, model card, demo recording",
        "Tight deadline — week 8 is non-negotiable. Prioritise deliverables over perfection"
    ),
]

for week_label, dates, phase, tasks, owner, deliverable, risk in weeks:
    # Week header row
    tbl_wk = doc.add_table(rows=5, cols=2)
    tbl_wk.style = 'Table Grid'

    # Row 0: Week + phase
    bg_col = NAVY if "1" in week_label or "2" in week_label else (
              PURPLE if "3" in week_label or "4" in week_label else (
              TEAL if "5" in week_label or "6" in week_label else
              AMBER if "7" in week_label else GREEN))

    tbl_wk.rows[0].cells[0].merge(tbl_wk.rows[0].cells[1])
    set_cell_bg(tbl_wk.rows[0].cells[0], bg_col)
    add_cell_para(tbl_wk.rows[0].cells[0],
                  f"{week_label}  ({dates})  —  {phase}",
                  bold=True, size=11.5, color=WHITE)

    # Row 1: Tasks
    set_cell_bg(tbl_wk.rows[1].cells[0], LGRAY)
    add_cell_para(tbl_wk.rows[1].cells[0], "Key Tasks", bold=True, size=10, color=NAVY)
    set_cell_bg(tbl_wk.rows[1].cells[1], LGRAY)
    tasks_text = "\n".join(f"  ▸  {t}" for t in tasks)
    tbl_wk.rows[1].cells[1].paragraphs[0].clear()
    run_t = tbl_wk.rows[1].cells[1].paragraphs[0].add_run(tasks_text)
    run_t.font.size = Pt(9.5); run_t.font.name = 'Calibri'
    run_t.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))

    # Row 2: Owner
    set_cell_bg(tbl_wk.rows[2].cells[0], WHITE)
    add_cell_para(tbl_wk.rows[2].cells[0], "Owner(s)", bold=True, size=10, color=NAVY)
    set_cell_bg(tbl_wk.rows[2].cells[1], WHITE)
    add_cell_para(tbl_wk.rows[2].cells[1], owner, size=10, color="1E293B")

    # Row 3: Deliverable
    set_cell_bg(tbl_wk.rows[3].cells[0], LGRAY)
    add_cell_para(tbl_wk.rows[3].cells[0], "Deliverable", bold=True, size=10, color=TEAL)
    set_cell_bg(tbl_wk.rows[3].cells[1], LGRAY)
    add_cell_para(tbl_wk.rows[3].cells[1], deliverable, size=10, color="1E293B")

    # Row 4: Risk
    set_cell_bg(tbl_wk.rows[4].cells[0], WHITE)
    add_cell_para(tbl_wk.rows[4].cells[0], "Risk / Mitigation", bold=True, size=10, color=RED)
    set_cell_bg(tbl_wk.rows[4].cells[1], WHITE)
    add_cell_para(tbl_wk.rows[4].cells[1], risk, size=10, italic=True, color=GRAY)

    para()

    # ── Orange vs Python callout (inserted after Week 2 block) ──────────────
    if "Week 2" in week_label:
        tbl_ov = doc.add_table(rows=2, cols=1)
        tbl_ov.style = 'Table Grid'
        set_cell_bg(tbl_ov.rows[0].cells[0], AMBER)
        add_cell_para(tbl_ov.rows[0].cells[0],
                      "TOOL DECISION NOTE — Why Python (scikit-learn / XGBoost) instead of Orange 3?",
                      bold=True, size=11, color=WHITE)
        set_cell_bg(tbl_ov.rows[1].cells[0], LGRAY)
        tbl_ov.rows[1].cells[0].paragraphs[0].clear()
        callout_body = (
            "Orange 3 was the visual ML tool taught during the APAL Programme at IIM Calcutta. "
            "It is excellent for exploring data, testing algorithms and explaining concepts without "
            "writing code. However, for this project we are using Python (scikit-learn + XGBoost + "
            "FastAPI) as the primary implementation language. The reasons are:\n\n"
            "  1.  Deployment requirement — Orange produces a visual workflow (.ows file) that cannot "
            "be packaged into an API or Docker container. Python .pkl model files can be loaded by "
            "FastAPI and served as a live HTTP endpoint that FinnUp's system calls in real-time.\n\n"
            "  2.  FinnUp integration — the final deliverable requires a REST API that FinnUp's "
            "portal calls when a borrower registers. This is only achievable in Python/FastAPI, "
            "not in Orange.\n\n"
            "  3.  Lender policy rule engine — applying per-lender eligibility filters (CIBIL "
            "minimums, sector constraints, geography) before ML scoring requires custom Python logic "
            "that Orange's widget-based environment cannot accommodate.\n\n"
            "  4.  SHAP explanations — the python-shap library integrates natively with XGBoost "
            "and scikit-learn pipelines, enabling automated 'top reasons' generation per prediction. "
            "Orange has limited SHAP support and no way to embed it in an API response.\n\n"
            "  5.  Batch scoring at scale — scoring all 6,483 existing borrowers in a nightly "
            "batch job (vectorized numpy/pandas operations) is orders of magnitude faster than "
            "Orange's row-by-row widget processing.\n\n"
            "  6.  Version control and reproducibility — Python code, model files and feature "
            "pipelines are tracked in Git and can be re-run by any team member. Orange workflows "
            "are tied to a GUI and difficult to version-control or automate.\n\n"
            "How Orange is still used (optional validation step):\n"
            "  > Orange is used optionally in Week 2 as a quick sanity-check tool — loading the "
            "feature matrix, visualising correlation matrices and running a quick Logistic Regression "
            "or Naive Bayes to confirm signals before committing to full XGBoost training in Python.\n"
            "  > Screenshots of Orange validation can be included in the IIM Calcutta final report "
            "as evidence of academic grounding alongside the production Python implementation.\n\n"
            "IIM Calcutta submission note:\n"
            "  'Initial model concept validated using Orange 3 (IIM APAL tool); production "
            "implementation built in Python (XGBoost / scikit-learn) to enable REST API deployment "
            "and live integration with FinnUp's borrower registration platform.'"
        )
        r_cb = tbl_ov.rows[1].cells[0].paragraphs[0].add_run(callout_body)
        r_cb.font.size = Pt(9.5); r_cb.font.name = 'Calibri'
        r_cb.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))
        para()

        ov_data = [
            ("Criteria",                    "Orange 3 (IIMC tool)",        "Python — scikit-learn + XGBoost"),
            ("Purpose",                     "Visual / teaching tool",      "Production ML system"),
            ("How you work",                "Drag-and-drop GUI widgets",   "Code in Jupyter notebooks + scripts"),
            ("Deployable as API?",          "No",                          "Yes — load .pkl into FastAPI"),
            ("FinnUp integration possible?","No",                          "Yes — core requirement"),
            ("SHAP explanations",           "Limited",                     "Full shap library support"),
            ("Lender rule engine",          "Not possible",                "Custom Python logic — full control"),
            ("Batch scoring 6,483 rows",    "Slow (widget row-by-row)",    "Fast — vectorised numpy/pandas"),
            ("Version control (Git)",       "Difficult (.ows binary)",     "Easy — .py and .pkl files"),
            ("IIM submission suitability",  "Classroom demo only",         "Professional deliverable + demo"),
        ]
        tbl_cmp = doc.add_table(rows=len(ov_data), cols=3)
        tbl_cmp.style = 'Table Grid'
        for r_i, (crit, orange_v, python_v) in enumerate(ov_data):
            row_bg = NAVY if r_i == 0 else (LGRAY if r_i % 2 == 0 else WHITE)
            txt_c  = WHITE if r_i == 0 else "1E293B"
            for c_i, txt in enumerate([crit, orange_v, python_v]):
                set_cell_bg(tbl_cmp.rows[r_i].cells[c_i], row_bg)
                is_bold = r_i == 0 or c_i == 0
                add_cell_para(tbl_cmp.rows[r_i].cells[c_i], txt,
                              bold=is_bold, size=9.5, color=txt_c)
        para()

divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 – SUMMARY TIMELINE TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 5 — TIMELINE SUMMARY AT A GLANCE", bg=TEAL)

tbl_sum = doc.add_table(rows=1 + 8, cols=5)
tbl_sum.style = 'Table Grid'
tbl_sum.alignment = WD_TABLE_ALIGNMENT.LEFT

for c, txt in zip(tbl_sum.rows[0].cells,
                  ["Week", "Dates", "Phase", "Key Output", "Status"]):
    set_cell_bg(c, NAVY)
    add_cell_para(c, txt, bold=True, size=10.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

summary_rows = [
    ("Jan–Mar 2026", "Jan–Mar", "Phase 0–1: Problem Definition + EDA", "EDA report, feature audit, label mapping", "✓ COMPLETE", GREEN),
    ("Week 1  (7–11 Apr)", "7–11 Apr", "Phase 2 Start: Labels + Data Augment", "Outcome labels; expanded dataset", "⏳ IMMEDIATE", AMBER),
    ("Week 2  (14–18 Apr)", "14–18 Apr", "Phase 2: Feature Engineering", "50+ feature matrix with imputation", "NOT STARTED", GRAY),
    ("Week 3  (21–25 Apr)", "21–25 Apr", "Phase 2: Labels + Train-Test Split", "Training-ready dataset; label series", "NOT STARTED", GRAY),
    ("Week 4  (28 Apr – 2 May)", "28 Apr", "Phase 3: Baseline + XGBoost Stage 1", "Trained model; AUC-ROC benchmark; SHAP", "NOT STARTED", GRAY),
    ("Week 5  (5–9 May)", "5–9 May", "Phase 3: Stage 2 Ranking + Model Card", "Ranked lender output; model card", "NOT STARTED", GRAY),
    ("Week 6  (12–16 May)", "12–16 May", "Phase 4: REST API", "Live FastAPI scoring endpoint", "NOT STARTED", GRAY),
    ("Week 7–8  (19–30 May)", "19–30 May", "Phase 4–5: Integration + Final Report", "UAT results; IIM final submission", "NOT STARTED", GRAY),
]

for i, (week, dates, phase, output, status, col) in enumerate(summary_rows):
    row = tbl_sum.rows[i + 1]
    bg3 = LGRAY if i % 2 == 0 else WHITE
    for j, (cell, txt2) in enumerate(zip(row.cells, [week, dates, phase, output, status])):
        set_cell_bg(cell, bg3)
        txt_color = col if j == 4 else "1E293B"
        add_cell_para(cell, txt2, size=9.5, color=txt_color, bold=(j == 4))

para()
divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — API ARCHITECTURE: HOW MODEL, API & FINNUP CONNECT
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 6 — API ARCHITECTURE: HOW THE MODEL, API & FINNUP CONNECT", bg=PURPLE)

heading("Overview", level=2)
para(
    "A common question: once the ML model is trained, how does it actually get used by FinnUp? "
    "This section explains the full flow end-to-end in plain language — from a borrower registering "
    "on the FinnUp platform to a lender receiving a ranked, scored lead.",
    space_after=6
)

heading("6.1  The Three Components (What Gets Built)", level=2)

components = [
    ("Component 1: The Trained ML Model  (a saved file)", NAVY, [
        "The model is trained once on historical FinnUp data (Weeks 4-5)",
        "After training it is saved as a file on disk — e.g. xgboost_model.pkl (a few MB in size)",
        "This file contains all learned patterns: which features predict approval for which lender",
        "It can be loaded instantly in memory and used to score thousands of borrowers per second",
        "Think of it like a trained brain saved to a USB drive — train once, use endlessly",
        "It gets retrained monthly as new approval/rejection outcomes arrive from FinnUp",
    ]),
    ("Component 2: The Feature Pipeline  (Python code that transforms raw data)", TEAL, [
        "Raw borrower data sent by FinnUp is messy — names, numbers, missing fields",
        "The feature pipeline transforms raw data into the 50+ numeric features the model expects",
        "Example: 'Company Type = Proprietorship' becomes a number the model understands (one-hot encode)",
        "Example: Turnover = Rs 5,00,000 becomes log(500000) = 13.1 (log transform)",
        "Example: if Director CIBIL is missing it fills it with the sector median and sets cibil_missing flag = 1",
        "This code runs automatically every time a borrower is scored — zero human involvement",
    ]),
    ("Component 3: The REST API  (the connector layer — built with FastAPI)", PURPLE, [
        "The API is a web service that runs on a cloud server 24/7 — like a website that never closes",
        "FinnUp sends a borrower's raw data to the API using a standard HTTP call",
        "The API runs the feature pipeline, feeds features to the model, applies lender rules, returns the score",
        "Response time: under 200 milliseconds — faster than a human eye blink",
        "Secured with an API key — only FinnUp's authorised system can call it",
        "Every call is audit-logged: timestamp, borrower ID, score returned",
    ]),
]

for title, bg, items in components:
    tbl_c = doc.add_table(rows=2, cols=1)
    tbl_c.style = 'Table Grid'
    set_cell_bg(tbl_c.rows[0].cells[0], bg)
    add_cell_para(tbl_c.rows[0].cells[0], title, bold=True, size=11, color=WHITE)
    set_cell_bg(tbl_c.rows[1].cells[0], LGRAY)
    body_text = "\n".join(f"  > {item}" for item in items)
    tbl_c.rows[1].cells[0].paragraphs[0].clear()
    r = tbl_c.rows[1].cells[0].paragraphs[0].add_run(body_text)
    r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))
    para()

divider()

heading("6.2  Step-by-Step: What Happens When a Borrower Registers", level=2)

para(
    "Here is exactly what happens from borrower registration to a lender seeing a scored lead "
    "— fully automated, no human needed at any step:",
    space_after=4
)

flow_steps = [
    ("STEP 1", NAVY,
     "Borrower submits registration on FinnUp portal",
     "A borrower fills in their details: company name, turnover, business type, location, loan amount needed.",
     "FinnUp's existing web portal — no changes needed to this step"),
    ("STEP 2", TEAL,
     "FinnUp's system sends the borrower data to our Scoring API",
     "Option A (Webhook / Real-time): FinnUp's system is configured to automatically call our "
     "API the moment a new borrower registers. Score is returned within 200ms.\n"
     "Option B (Batch / Nightly): A scheduled script runs at midnight, collects all new "
     "borrowers registered that day, sends them all to the API in bulk, writes scores back to FinnUp's CRM.",
     "FinnUp Tech adds 1 line: POST https://api.finnup-score.com/v1/score  (1-day integration task)"),
    ("STEP 3", PURPLE,
     "API receives raw borrower data as a JSON message and validates it",
     'Example JSON received by API:\n'
     '  { "borrower_id": "B12345", "turnover": 5000000, "company_type": "Proprietorship",\n'
     '    "state": "Maharashtra", "director_cibil": 720, "loan_amount": 2000000 }\n\n'
     "API validates all fields (required fields, data types, value ranges) before processing anything.",
     "Input validation handled automatically by Pydantic library — invalid data rejected with error message"),
    ("STEP 4", AMBER,
     "Feature pipeline transforms raw data into 50+ model-ready features (takes ~5ms)",
     "  > log(5000000) = 15.4  (log-transform of Turnover)\n"
     "  > company_type_proprietorship = 1  (one-hot encoding)\n"
     "  > state_maharashtra = 1  (state encoding)\n"
     "  > director_cibil = 720,  director_cibil_imputed = 0  (real value present)\n"
     "  > has_bank_data = 0,  bounce_ratio = 0  (no bank data yet — flagged)\n"
     "  > ... 50+ features computed in total",
     "Same feature pipeline code built in Weeks 2-3; reused here inside the API"),
    ("STEP 5", RED,
     "Lender policy rule engine filters eligible lenders (hard rules applied first)",
     "Before ML runs, the rule engine checks which lenders this borrower even qualifies for:\n"
     "  > Lender A: needs CIBIL >= 680, Turnover >= Rs 30L, Maharashtra only  -> PASS\n"
     "  > Lender B: needs CIBIL >= 750, Manufacturing sector only  -> FAIL (CIBIL only 720)\n"
     "  > Lender C: no CIBIL minimum, Turnover >= Rs 10L  -> PASS\n"
     "Result: 4 of 10 lenders are eligible. ML only runs on these 4 — saves compute, ensures compliance.",
     "Lender policy JSON files already in /data/lender_policies/ — confirmed in Week 1"),
    ("STEP 6", TEAL,
     "XGBoost model predicts P(approval) for each eligible lender and ranks them",
     "  > Lender A + Borrower B12345  ->  P(approval) = 0.82  ->  Score = 820\n"
     "  > Lender C + Borrower B12345  ->  P(approval) = 0.71  ->  Score = 710\n"
     "  > Lender D + Borrower B12345  ->  P(approval) = 0.61  ->  Score = 610\n"
     "  > Lender E + Borrower B12345  ->  P(approval) = 0.45  ->  Score = 450\n\n"
     "Lenders are ranked: A (820) > C (710) > D (610) > E (450)\n"
     "Top 3 recommended: Lender A, C, D",
     "XGBoost model trained and saved as .pkl file in Week 4; loaded into memory when API starts"),
    ("STEP 7", GREEN,
     "SHAP generates human-readable explanation for WHY this score was given",
     "SHAP breaks the prediction into each feature's contribution. Example output:\n"
     "  > +120 points: Director CIBIL = 720  (above-average credit signal — positive)\n"
     "  > +85 points:  Turnover Rs 50L  (good scale for this loan amount — positive)\n"
     "  > -40 points:  No bank statement data available  (uncertainty — negative)\n"
     "  > -25 points:  Business Age < 2 years  (young business — slight negative)\n\n"
     "These become the 'Top reasons for this score' section on the lender's screen.",
     "SHAP library integrated directly into the API scoring function — runs automatically"),
    ("STEP 8", NAVY,
     "API sends the complete scorecard back to FinnUp as a JSON response (within 200ms)",
     '  {\n'
     '    "borrower_id": "B12345",\n'
     '    "credit_score": 820,  "risk_band": "Good",\n'
     '    "top_lender_matches": [\n'
     '      { "lender": "Lender A", "score": 820, "product": "Term Loan" },\n'
     '      { "lender": "Lender C", "score": 710, "product": "Working Capital" }\n'
     '    ],\n'
     '    "key_positive_factors": ["Director CIBIL 720", "Turnover Rs 50L"],\n'
     '    "key_negative_factors": ["No bank data", "Business age < 2 years"],\n'
     '    "data_quality_flags": ["bank_data_missing"]\n'
     '  }',
     "Standard JSON format — FinnUp reads this and updates the borrower lead record in their CRM"),
    ("STEP 9", GREEN,
     "FinnUp's dashboard shows the scored lead to the matched lender",
     "The lender logs into the FinnUp portal and sees:\n"
     "  Credit Score: 820  |  Risk Band: Good\n"
     "  Recommended action: High priority — contact within 24 hours\n"
     "  Top positive factors: Director CIBIL 720, Strong Turnover\n\n"
     "Optionally: FinnUp sends an email or WhatsApp alert to the lender's relationship manager.",
     "FinnUp reads the JSON response from Step 8 and renders it in their existing dashboard"),
]

for step_idx, (step_label, bg, step_title, detail, tech_note) in enumerate(flow_steps):
    tbl_s = doc.add_table(rows=1, cols=2)
    tbl_s.style = 'Table Grid'
    tbl_s.rows[0].cells[0].width = Inches(0.9)
    set_cell_bg(tbl_s.rows[0].cells[0], bg)
    add_cell_para(tbl_s.rows[0].cells[0], step_label, bold=True, size=11, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    content_col = tbl_s.rows[0].cells[1]
    set_cell_bg(content_col, LGRAY if step_idx % 2 == 0 else WHITE)
    content_col.paragraphs[0].clear()
    r_title = content_col.paragraphs[0].add_run(step_title + "\n")
    r_title.bold = True; r_title.font.size = Pt(10.5)
    r_title.font.name = 'Calibri'
    r_title.font.color.rgb = RGBColor(*hex_to_rgb(bg))
    r_detail = content_col.paragraphs[0].add_run(detail + "\n")
    r_detail.font.size = Pt(9.5); r_detail.font.name = 'Calibri'
    r_detail.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))
    r_note = content_col.paragraphs[0].add_run("  Tech note: " + tech_note)
    r_note.font.size = Pt(9); r_note.font.name = 'Calibri'; r_note.font.italic = True
    r_note.font.color.rgb = RGBColor(*hex_to_rgb(GRAY))
    para()

divider()

heading("6.3  The Two Integration Options for FinnUp", level=2)
para("FinnUp can use either mode — or both:", space_after=4)

tbl_int = doc.add_table(rows=3, cols=4)
tbl_int.style = 'Table Grid'
for c, txt in zip(tbl_int.rows[0].cells, ["Mode", "How It Works", "Best For", "FinnUp Setup Needed"]):
    set_cell_bg(c, NAVY)
    add_cell_para(c, txt, bold=True, size=10.5, color=WHITE)

int_modes = [
    ("Option A — Webhook (Real-Time)",
     "FinnUp adds one config line to their registration form. When a borrower submits, "
     "FinnUp automatically POSTs the data to our API URL. Score is returned in 200ms. "
     "FinnUp stores it and can show it immediately on the borrower confirmation page.",
     "New borrower onboarding. Score available before the registration confirmation is even shown to the borrower.",
     "FinnUp Tech configures 1 webhook URL in their form.\nEstimated effort: 1 day"),
    ("Option B — Nightly Batch Job",
     "A scheduled script runs at midnight. It queries FinnUp's DB for all new "
     "borrowers that day, sends them to POST /v1/score/batch, receives all scores, "
     "and writes them back to FinnUp's CRM database.",
     "Scoring the entire existing 6,483 borrower base in one go. Re-scoring all borrowers "
     "when lender policies change. No latency requirement.",
     "Script runs on cloud (cron job).\nFinnUp grants DB read-access.\nEstimated effort: 2-3 days"),
]

for i, (mode, how, best, setup) in enumerate(int_modes):
    row = tbl_int.rows[i + 1]
    bg_i = LGRAY if i == 0 else WHITE
    for c, txt in zip(row.cells, [mode, how, best, setup]):
        set_cell_bg(c, bg_i)
        add_cell_para(c, txt, size=9.5, color="1E293B", bold=(i == 0 and c == row.cells[0]))

para()
divider()

heading("6.4  Weeks 5, 6 and 7 Explained Simply", level=2)

week_detail_rows = [
    ("WEEK 5  (5-9 May) — Stage 2 Lender Ranking + Model Card", PURPLE,
     "What is Stage 2?\n"
     "Week 4's XGBoost model outputs P(approval) per borrower. But FinnUp has 10+ lenders — "
     "which specific ones do we recommend to THIS borrower? Stage 2 answers that.\n\n"
     "How it works (3 steps):\n"
     "  Step A: Run lender policy rules — filter out lenders whose hard requirements fail\n"
     "  Step B: For each remaining eligible lender, run XGBoost to get their P(approval)\n"
     "  Step C: Rank eligible lenders from highest to lowest score — top 3 become the recommendation\n\n"
     "Example:\n"
     "  Borrower: Director CIBIL 720, Turnover Rs 40L\n"
     "  8 of 10 lenders pass the hard rules filter\n"
     "  XGBoost scores them: A=0.82, E=0.71, C=0.65, D=0.55...\n"
     "  Final output: 'Top 3 recommended: Lender A, E, C'\n\n"
     "What is the Model Card?\n"
     "A 2-page document (like a product datasheet) formally describing the model:\n"
     "  > What it predicts and what data it was trained on\n"
     "  > Performance metrics: AUC-ROC, Precision, Recall on the test set\n"
     "  > Known limitations: small dataset size, proxy labels used instead of true outcomes\n"
     "  > Fairness analysis: does the model score differently by state or company type?\n"
     "  > Intended use warnings (what the model should and should not be used for)\n"
     "Required by IIM Calcutta for the final submission and by FinnUp for compliance records."),

    ("WEEK 6  (12-16 May) — REST API Development", NAVY,
     "What is FastAPI?\n"
     "FastAPI is a Python framework for building web services (like Django builds websites, "
     "FastAPI builds APIs). All the scoring logic — feature pipeline + model + lender ranking — "
     "is wrapped inside this API so FinnUp's system can call it without knowing the internal code.\n\n"
     "What endpoints are built?\n"
     "  POST /v1/score\n"
     "    Send: one borrower's raw data as JSON\n"
     "    Receive: credit score, risk band, top lenders, SHAP explanation\n"
     "    Used by: FinnUp webhook (real-time, one borrower at a time)\n\n"
     "  POST /v1/score/batch\n"
     "    Send: list of borrowers as JSON array\n"
     "    Receive: scores for all of them in one call\n"
     "    Used by: nightly batch job\n\n"
     "  GET /v1/lenders/match?borrower_id=B12345\n"
     "    Returns: cached lender match list for that borrower\n"
     "    Used by: FinnUp dashboard to display the recommendation\n\n"
     "  GET /v1/health\n"
     "    Returns: API status, model version, last updated\n"
     "    Used by: monitoring to confirm API is alive\n\n"
     "How is security handled?\n"
     "  > API Key: FinnUp includes a secret key in every request — invalid keys rejected\n"
     "  > Input validation: Pydantic checks every field before it reaches the model\n"
     "  > HTTPS only: all traffic encrypted in transit\n"
     "  > Audit log: every call logged with timestamp and borrower ID (no PII in logs)\n\n"
     "What is Docker?\n"
     "Docker packages the entire application (code + model file + dependencies) into a "
     "self-contained container that runs identically on any computer or cloud server. "
     "This means: no 'works on my machine' issues, easy to deploy to AWS or Azure with "
     "one command, and easy to update the model without downtime."),

    ("WEEK 7  (19-23 May) — FinnUp Integration & UAT", TEAL,
     "Deployment to cloud:\n"
     "The Docker container from Week 6 is deployed to a cloud server (AWS or Azure). "
     "This gives the API a real public HTTPS URL (e.g. https://api.finnup-scoring.com). "
     "The server is always on, auto-scales, and has the model loaded in memory.\n\n"
     "FinnUp integration — 3 steps (done jointly with FinnUp Tech team):\n"
     "  Step 1: We share the API URL + FinnUp's API key with their tech team\n"
     "  Step 2: FinnUp's developer adds the POST call to their registration form (1-day task)\n"
     "  Step 3: FinnUp maps credit_score, risk_band, top_lender_matches from API response "
     "into fields in their borrower database for display on the dashboard\n\n"
     "What is UAT (User Acceptance Testing)?\n"
     "After the API is live and connected, we run a formal test with FinnUp's team:\n"
     "  > Take 200 real FinnUp borrowers whose outcomes are known\n"
     "  > Send them through the API and review the scores + lender recommendations\n"
     "  > FinnUp's team checks: do the recommended lenders make sense for these borrowers?\n"
     "  > Are the SHAP explanations understandable to a loan officer?\n"
     "  > Is the scorecard format clear on the lender dashboard?\n"
     "  > Are there edge cases where the API fails or gives unexpected results?\n"
     "UAT sign-off from FinnUp is required before calling the system production-ready."),
]

for wk_title, bg, body in week_detail_rows:
    tbl_wd = doc.add_table(rows=2, cols=1)
    tbl_wd.style = 'Table Grid'
    set_cell_bg(tbl_wd.rows[0].cells[0], bg)
    add_cell_para(tbl_wd.rows[0].cells[0], wk_title, bold=True, size=12, color=WHITE)
    set_cell_bg(tbl_wd.rows[1].cells[0], LGRAY)
    tbl_wd.rows[1].cells[0].paragraphs[0].clear()
    r_b = tbl_wd.rows[1].cells[0].paragraphs[0].add_run(body)
    r_b.font.size = Pt(10); r_b.font.name = 'Calibri'
    r_b.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))
    para()

divider()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — FINAL DELIVERABLES FOR IIM CALCUTTA
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 7 — FINAL DELIVERABLES FOR IIM CALCUTTA", bg=NAVY)

heading("What will be submitted by 31 May 2026:", level=2)

deliverables = [
    ("Project Report (this document, updated)",
     "Full write-up of problem, methodology, data, model, results, and business impact"),
    ("EDA Report (PowerPoint)",
     "Auto-generated EDA deck (FinnUp_EDA_Report.pptx) — already complete"),
    ("Feature Engineering Notebook",
     "02_feature_engineering.ipynb — documented feature pipeline with rationale"),
    ("Model Training Notebook",
     "03_model_training.ipynb — XGBoost model with SHAP, AUC-ROC charts, validation results"),
    ("Model Card",
     "Formal model card: performance, fairness audit, known limitations, intended use"),
    ("API Documentation",
     "OpenAPI/Swagger documentation for the REST scoring endpoint"),
    ("Business Impact Report",
     "Measured or projected improvement over rules-based system (A/B comparison)"),
    ("Final Presentation (PowerPoint)",
     "Updated from this interim submission — includes model results and live demo screenshots"),
    ("GitHub / Code Repository",
     "Complete codebase: loader, feature engineering, model training, API — all documented"),
    ("Demo Recording (optional)",
     "Short video of the live scoring API returning results for a sample borrower"),
]

for title, detail in deliverables:
    bullet(f"{title} — {detail}", bold_prefix="◉", prefix_color=TEAL)

para()
divider()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 – OPEN BLOCKERS & ACTIONS
# ══════════════════════════════════════════════════════════════════════════════
add_banner("SECTION 8 — OPEN BLOCKERS & REQUIRED ACTIONS", bg=RED)

heading("Critical items requiring immediate action (Week 1):", level=2)

blockers = [
    ("🔴 Outcome Labels from FinnUp CRM", "CRITICAL BLOCKER",
     "Request approved/rejected/disbursed status per borrower-loan pair. "
     "Export from LOS/CRM team. Even 1,000 labelled records unblock Phase 3. "
     "FinnUp Status (YES/NO) will serve as interim proxy."),
    ("🔴 Lender Policy File Confirmation", "CRITICAL BLOCKER",
     "Confirm /data/lender_policies/ is complete and up to date. "
     "Each lender's minimum CIBIL, turnover floor, sector whitelist, and ticket range is needed."),
    ("🟡 More Borrower Records", "IMPORTANT",
     "Request all new registrations since original export. Also request historical data "
     "going back 3+ years. Target: 10,000+ borrower records for reliable model training."),
    ("🟡 Bank Statement / AA Coverage", "IMPORTANT",
     "Confirm whether Account Aggregator consent data is available for the 62% uncovered. "
     "Even a push to 50% bank coverage significantly improves model quality."),
    ("⚪ FinnUp Tech Team Contact", "NICE TO HAVE",
     "Introduce FinnUp engineering team for Phase 4 integration (webhook + API deployment). "
     "Schedule kickoff by Week 5."),
]

tbl_bl = doc.add_table(rows=1 + len(blockers), cols=3)
tbl_bl.style = 'Table Grid'
for c, txt in zip(tbl_bl.rows[0].cells, ["Item", "Priority", "Detail"]):
    set_cell_bg(c, RED)
    add_cell_para(c, txt, bold=True, size=10.5, color=WHITE)
for i, (item, pri, detail) in enumerate(blockers):
    row = tbl_bl.rows[i + 1]
    bg4 = LGRAY if i % 2 == 0 else WHITE
    set_cell_bg(row.cells[0], bg4); add_cell_para(row.cells[0], item, bold=True, size=10, color=NAVY)
    pcol = RED if "CRITICAL" in pri else (AMBER if "IMPORTANT" in pri else GRAY)
    set_cell_bg(row.cells[1], bg4); add_cell_para(row.cells[1], pri, bold=True, size=10, color=pcol)
    set_cell_bg(row.cells[2], bg4); add_cell_para(row.cells[2], detail, size=10, color="1E293B")

para()
divider()


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

p_close = doc.add_paragraph()
p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_close.paragraph_format.space_before = Pt(40)
rc = p_close.add_run("FinnUp MSME Lender Matching — Project Summary")
rc.bold = True; rc.font.size = Pt(16); rc.font.name = 'Calibri'
rc.font.color.rgb = RGBColor(*hex_to_rgb(NAVY))

closing_bullets = [
    ("✓  DONE", GREEN,   "Phases 0–1 complete: 6,483 borrowers, 11 sheets, 30+ ML features, AI Canvas validated"),
    ("✓  DONE", GREEN,   "EDA delivered: missing value map, feature coverage, label landscape, bank risk signals"),
    ("⏳  NOW",  AMBER,   "Week 1: Collect outcome labels from FinnUp — this is the critical dependency"),
    ("◌  Apr",  PURPLE,  "Weeks 2–3: Feature engineering, imputation, and training dataset creation"),
    ("◌  May",  TEAL,    "Weeks 4–5: XGBoost model training, SHAP explanations, model card"),
    ("◌  May",  NAVY,    "Weeks 6–7: REST API development, FinnUp integration, UAT"),
    ("◌  May",  GREEN,   "Week 8: Final report, IIM Calcutta submission, stakeholder demo — DEADLINE: 31 May 2026"),
]

p_sum = doc.add_paragraph()
p_sum.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sum.paragraph_format.space_before = Pt(12)

for badge, col, text in closing_bullets:
    p_b = doc.add_paragraph()
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_b.add_run(f"{badge}  ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r1.font.color.rgb = RGBColor(*hex_to_rgb(col))
    r2 = p_b.add_run(text)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'
    r2.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))
    p_b.paragraph_format.space_after = Pt(3)

para()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_foot.add_run(
    "APAL Programme  |  IIM Calcutta  |  Cohort 2  |  Group 1  |  April 2026  |  Confidential"
)
rf.font.size = Pt(9); rf.font.name = 'Calibri'
rf.font.color.rgb = RGBColor(*hex_to_rgb(GRAY))


# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"\n✓  Saved: {OUTPUT_PATH.resolve()}")
