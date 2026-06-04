"""
generate_iim_summary_report.py
-------------------------------
Generates  outputs/FinnUp_IIM_Summary_Report.docx
  – APAL Cohort 2 | Group 1 | IIM Calcutta
    Compact 5-page summary: Status + Plan + Architecture + Deliverables

Run:
    python generate_iim_summary_report.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = Path("outputs") / "FinnUp_IIM_Summary_Report.docx"
OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

# ── Colours ────────────────────────────────────────────────────────────────────
def hex_to_rgb(h): return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

NAVY   = "1B3A6B"
TEAL   = "0D9488"
GREEN  = "16A34A"
AMBER  = "F59E0B"
RED    = "EF4444"
PURPLE = "7C3AED"
GRAY   = "64748B"
WHITE  = "FFFFFF"
LGRAY  = "F5F7FA"
LGREEN = "DCFCE7"
LAMBER = "FEF3C7"
LRED   = "FEE2E2"

ALL_MEMBERS = (
    "Asha, Arvind, Anil, Bhupesh, Deepak, Ganesh, "
    "Gopal, Hareram, Pranali, Rahul, Savitha, Sonam, Samik"
)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_cell_para(cell, text, bold=False, size=10, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    return p

# ── Document ───────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.page_height  = Cm(29.7)
    section.page_width   = Cm(21.0)
    section.top_margin   = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def p(text='', bold=False, italic=False, color=None, size=10.5, space_after=3):
    pg = doc.add_paragraph()
    pg.paragraph_format.space_before = Pt(0)
    pg.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = pg.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = 'Calibri'
        if color:
            r.font.color.rgb = RGBColor(*hex_to_rgb(color))
    return pg

def banner(text, bg=NAVY, size=12.5):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    add_cell_para(cell, text, bold=True, size=size, color=WHITE)
    p()

def mini_heading(text, color=NAVY):
    pg = doc.add_paragraph()
    pg.paragraph_format.space_before = Pt(6)
    pg.paragraph_format.space_after  = Pt(2)
    r = pg.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(*hex_to_rgb(color))

def divider():
    pg = doc.add_paragraph()
    pg.paragraph_format.space_before = Pt(0)
    pg.paragraph_format.space_after  = Pt(0)
    r = pg.add_run('─' * 105)
    r.font.size = Pt(7); r.font.color.rgb = RGBColor(*hex_to_rgb(GRAY))

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — TITLE + EXECUTIVE SUMMARY + PROJECT STATUS
# ══════════════════════════════════════════════════════════════════════════════

# ── Title block ────────────────────────────────────────────────────────────────
tbl_title = doc.add_table(rows=1, cols=1)
tbl_title.alignment = WD_TABLE_ALIGNMENT.LEFT
tc = tbl_title.rows[0].cells[0]
set_cell_bg(tc, NAVY)
tc.paragraphs[0].clear()
tc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = tc.paragraphs[0].add_run(
    "FinnUp MSME Credit Intelligence Platform\n"
    "AI-Powered Lender Matching — Project Summary\n"
)
r1.bold = True; r1.font.size = Pt(15); r1.font.name = 'Calibri'
r1.font.color.rgb = RGBColor(*hex_to_rgb(WHITE))
r2 = tc.paragraphs[0].add_run(
    "APAL Programme  |  IIM Calcutta  |  Cohort 2 – Group 1  |  April 2026"
)
r2.font.size = Pt(10); r2.font.name = 'Calibri'
r2.font.color.rgb = RGBColor(*hex_to_rgb("DBEAFE"))
p()

# ── Team strip ─────────────────────────────────────────────────────────────────
tbl_team = doc.add_table(rows=1, cols=1)
tc2 = tbl_team.rows[0].cells[0]
set_cell_bg(tc2, TEAL)
tc2.paragraphs[0].clear()
tc2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = tc2.paragraphs[0].add_run(f"Team (13 members):  {ALL_MEMBERS}")
rt.font.size = Pt(9.5); rt.font.name = 'Calibri'
rt.font.color.rgb = RGBColor(*hex_to_rgb(WHITE))
p()

# ── Executive summary 3-col table: Done / In Progress / Planned ───────────────
mini_heading("Executive Summary")
tbl_exec = doc.add_table(rows=2, cols=3)
tbl_exec.style = 'Table Grid'
for c, (hdr, bg) in enumerate(zip(
        ["✓  COMPLETED", "⏳  IMMEDIATE (Week 1)", "◌  PLANNED (Weeks 2–8)"],
        [GREEN, AMBER, NAVY])):
    set_cell_bg(tbl_exec.rows[0].cells[c], bg)
    add_cell_para(tbl_exec.rows[0].cells[c], hdr, bold=True, size=10, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

done_txt = (
    "Phase 0: Problem scoping\n"
    "Phase 1: Full EDA (11 sections)\n"
    "Data extraction from FinnUp DB\n"
    "6,483 borrowers across 11 Excel sheets\n"
    "Feature audit: 50+ ML-ready signals\n"
    "Label mapping: FinnUp Status (Y/N)\n"
    "Director CIBIL: 86% coverage confirmed\n"
    "Bank data: 38% coverage mapped\n"
    "Missing value strategy defined"
)
now_txt = (
    "Collect outcome labels from FinnUp\n"
    "Expand training rows (multi-year)\n"
    "Confirm lender policy files\n"
    "Clarify AA/bank consent (62% gap)\n"
    "All 13 members aligned on Week 2 scope"
)
plan_txt = (
    "Feature engineering (50+ features)\n"
    "XGBoost + LightGBM training\n"
    "Stage 2 lender ranking model\n"
    "SHAP-based explanations\n"
    "FastAPI REST scoring endpoint\n"
    "Docker containerisation\n"
    "FinnUp integration + UAT\n"
    "IIM final submission (May 31)"
)

for c, txt, bg in zip(range(3), [done_txt, now_txt, plan_txt], [LGREEN, LAMBER, LGRAY]):
    set_cell_bg(tbl_exec.rows[1].cells[c], bg)
    tbl_exec.rows[1].cells[c].paragraphs[0].clear()
    r = tbl_exec.rows[1].cells[c].paragraphs[0].add_run(txt)
    r.font.size = Pt(9); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))

p()
divider()

# ── Key data facts strip ────────────────────────────────────────────────────────
mini_heading("Key Data Facts (from EDA)")
facts = [
    ("6,483", "Unique borrowers\n(FinnUp DB)"),
    ("86%", "Director CIBIL\ncoverage"),
    ("38%", "Bank statement\ncoverage"),
    ("0.8%", "Borrower CIBIL\ncoverage"),
    ("50+", "ML-ready\nfeatures"),
    ("May 31", "Target\ncompletion"),
]
tbl_facts = doc.add_table(rows=2, cols=len(facts))
tbl_facts.style = 'Table Grid'
for i, (num, lbl) in enumerate(facts):
    bg_f = [NAVY, TEAL, PURPLE, AMBER, GREEN, RED][i]
    set_cell_bg(tbl_facts.rows[0].cells[i], bg_f)
    add_cell_para(tbl_facts.rows[0].cells[i], num, bold=True, size=15, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl_facts.rows[1].cells[i], LGRAY)
    add_cell_para(tbl_facts.rows[1].cells[i], lbl, size=8.5, color=GRAY,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
p()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — WHAT HAS BEEN COMPLETED (EDA FINDINGS)
# ══════════════════════════════════════════════════════════════════════════════
banner("PAGE 2 — WORK COMPLETED: PHASE 0 & PHASE 1 (EDA)", bg=GREEN)

mini_heading("Phase 0: Problem Definition & Setup (Jan–Feb 2026)")
completed_p0 = [
    "Defined business problem: MSME borrowers on FinnUp platform are unmatched to lenders due to incomplete credit profiles",
    "Selected AI Canvas framework to validate problem-solution fit across all 6 canvas stages",
    "Confirmed data source: FinnUp production database extract → FinnUp_Borrowers.xlsx (11 sheets)",
    "Mapped available labels: FinnUp Status (YES = matched / NO = not matched) as binary training target",
    "Confirmed two-stage ML architecture: Stage 1 = credit scoring, Stage 2 = lender ranking",
    "Identified IIM Calcutta submission requirements and academic framework (explainability, fairness, model card)",
]
for item in completed_p0:
    pg = doc.add_paragraph(style='List Bullet')
    pg.paragraph_format.space_after = Pt(1)
    pg.paragraph_format.left_indent = Inches(0.3)
    r = pg.add_run(item)
    r.font.size = Pt(10); r.font.name = 'Calibri'

p()
mini_heading("Phase 1: Full EDA — 11 Sections Completed (Feb–Mar 2026)")

eda_findings = [
    ("Data Coverage Analysis",
     "Borrower CIBIL: only 55/6,483 (0.8%) — unusable as primary feature. "
     "Director CIBIL: 5,587/6,483 (86%) — confirmed as primary credit proxy. "
     "Bank data: 2,489/6,483 (38%) — valuable secondary signal where available."),
    ("Company Type Distribution",
     "Proprietorships dominate (67%). Private Ltd = 18%, Partnership = 9%, Others = 6%. "
     "Company type is a strong feature — encodes risk profile and lender eligibility."),
    ("Director CIBIL Analysis",
     "Mean CIBIL 703; range 300–900. CIBIL > 700 correlates strongly with FinnUp YES status. "
     "14% missing — will be imputed using sector-median regression in feature engineering."),
    ("Bank Statement Risk Signals",
     "Bounce ratio (bounces/total transactions), CDR (credit-debit ratio), and EOD balance mean/std "
     "quantified for 2,489 borrowers with bank data. High bounce ratio = strong negative signal."),
    ("Financial KPIs",
     "Turnover 100% coverage — primary financial scale feature. Networth 92% coverage. "
     "DSCR (Debt Service Coverage Ratio) computed where loan + financial data available."),
    ("Geographic Distribution",
     "Maharashtra (34%), Karnataka (18%), Delhi NCR (11%) are top states. "
     "State encoding needed — both one-hot and target encoding will be tested."),
    ("Loan Product Analysis",
     "Term Loan 58%, Working Capital 27%, ODCC 9%, Invoice Discounting 6%. "
     "Loan amount relative to turnover (loan-to-turnover ratio) is a key derived feature."),
    ("Correlation & Feature Importance",
     "Top correlated features: Director CIBIL, Turnover, Networth, Bounce Ratio, Business Age. "
     "30+ features confirmed ML-ready; feature importance analysis ready for Week 4."),
    ("Missing Value Strategy",
     "Director CIBIL missing → sector-median imputation + cibil_missing flag. "
     "Bank data missing → fill 0 + has_bank_data flag. "
     "Financial KPIs missing → sector-median + kpi_missing flag."),
    ("Label Quality Assessment",
     "FinnUp Status: 62% YES, 38% NO — mild imbalance, manageable with SMOTE or class weights. "
     "Confirmed as best available proxy label; true loan outcome labels to be collected in Week 1."),
    ("Data Quality Summary",
     "11 sheets joined on Borrower ID. Key join: Profile + Director + Bank (38% overlap). "
     "No data leakage risk confirmed — all features are pre-application borrower attributes."),
]

tbl_eda = doc.add_table(rows=1 + len(eda_findings), cols=2)
tbl_eda.style = 'Table Grid'
for c, hdr in zip(tbl_eda.rows[0].cells, ["EDA Section", "Key Finding"]):
    set_cell_bg(c, TEAL)
    add_cell_para(c, hdr, bold=True, size=10, color=WHITE)

for i, (section, finding) in enumerate(eda_findings):
    row = tbl_eda.rows[i + 1]
    bg_e = LGRAY if i % 2 == 0 else WHITE
    set_cell_bg(row.cells[0], bg_e)
    add_cell_para(row.cells[0], section, bold=True, size=9.5, color=NAVY)
    set_cell_bg(row.cells[1], bg_e)
    add_cell_para(row.cells[1], finding, size=9.5, color="1E293B")

p()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — WEEK-BY-WEEK PLAN (ALL 13 MEMBERS ASSIGNED)
# ══════════════════════════════════════════════════════════════════════════════
banner("PAGE 3 — WEEK-BY-WEEK PLAN TO MAY 31, 2026 (ALL 13 MEMBERS)", bg=NAVY)

p("All 13 team members are assigned across every week. "
  "Week 8 (final report + submission) involves the full team.", size=9.5, italic=True, color=GRAY)
p()

weeks_plan = [
    (
        "Week 1", "7–11 Apr", NAVY,
        "Phase 2 Start: Labels + Data Augmentation",
        [
            "Collect outcome labels (approved/rejected) from FinnUp for borrowers with known outcomes",
            "Validate and lock FinnUp Status (Y/N) as proxy label — data agreement with FinnUp team",
            "Explore multi-year historical loan applications to expand training dataset",
            "Confirm all 11 lender policy files are complete, current and parsed correctly",
            "Clarify AA (Account Aggregator) consent status — plan for 62% uncovered borrowers",
            "Set up shared Git repository; all members aligned on folder structure and notebook naming",
        ],
        "Ganesh, Hareram, Bhupesh, Savitha (data leads) + Asha, Arvind (policy review)",
        "Remaining members: Anil, Deepak, Gopal, Pranali, Rahul, Sonam, Samik (environment setup + EDA review)"
    ),
    (
        "Week 2", "14–18 Apr", TEAL,
        "Phase 2: Feature Engineering",
        [
            "Build notebook 02_feature_engineering.ipynb — all transforms documented",
            "Engineer Profile features: log(Turnover), log(LoanAmt), business age, company type (one-hot)",
            "Engineer Director CIBIL features: min/max/avg across directors, cibil_missing flag",
            "Build CIBIL imputation sub-model (regress on sector + financial features)",
            "Engineer Bank Statement features: bounce_ratio, CDR, EOD_balance_mean, EOD_balance_std",
            "Implement all missing value strategies identified in EDA (sector-median + flags)",
        ],
        "Pranali, Rahul, Anil, Asha (feature engineering leads)",
        "Remaining members: Arvind, Bhupesh, Deepak, Ganesh, Gopal, Hareram, Savitha, Sonam, Samik (review + validation)"
    ),
    (
        "Week 3", "21–25 Apr", PURPLE,
        "Phase 2 cont.: Labels + Train-Test Split + Feature Selection",
        [
            "Engineer Loan Request features: product type encoding, loan-to-turnover ratio",
            "Engineer Financial KPI features: DSCR, networth-to-loan ratio, sector-median imputation",
            "Build all derived / cross-sheet features (e.g. director_avg_cibil × has_bank_data)",
            "Create final binary training label (1 = approved/matched, 0 = not)",
            "Train-test split: 80/20 stratified by label; temporal holdout for last-3-month test set",
            "Feature selection: remove near-zero variance and highly correlated pairs (r > 0.95)",
        ],
        "Pranali, Rahul, Deepak, Sonam (labels + feature selection)",
        "Remaining members: Asha, Anil, Arvind, Bhupesh, Ganesh, Gopal, Hareram, Savitha, Samik (peer review)"
    ),
    (
        "Week 4", "28 Apr – 2 May", PURPLE,
        "Phase 3: Baseline + XGBoost Stage 1 Model Training",
        [
            "Train Logistic Regression baseline (top-10 features) — establish AUC-ROC benchmark",
            "Train XGBoost Stage 1 classifier: predict P(approval) per borrower-lender pair",
            "Apply SMOTE oversampling + scale_pos_weight to handle 62/38 class imbalance",
            "Hyperparameter tuning: Optuna / RandomizedSearchCV with 5-fold stratified CV",
            "Evaluate: AUC-ROC, Precision, Recall, F1, Confusion Matrix on holdout test set",
            "SHAP feature importance: identify top-15 predictive features across all lenders",
        ],
        "Rahul, Ganesh, Arvind, Gopal (model training leads)",
        "Remaining members: Asha, Anil, Bhupesh, Deepak, Hareram, Pranali, Savitha, Sonam, Samik (results review)"
    ),
    (
        "Week 5", "5–9 May", TEAL,
        "Phase 3 cont.: Stage 2 Lender Ranking + Model Card",
        [
            "Build Stage 2: lender policy rule engine (hard eligibility filter — CIBIL min, sector, state)",
            "For each eligible lender: run XGBoost to get P(approval); rank lenders highest to lowest",
            "Define credit score bands: Poor (<400), Fair (400–600), Good (600–750), Excellent (750+)",
            "Generate SHAP-based 'Top 5 reasons' explanation per prediction",
            "Write Model Card: metrics, limitations, fairness analysis by state and company type",
            "End-to-end test on 200 sample borrowers — validate ranked output matches intuition",
        ],
        "Rahul, Gopal, Samik, Savitha (lender ranking + model card)",
        "Remaining members: Asha, Anil, Arvind, Bhupesh, Deepak, Ganesh, Hareram, Pranali, Sonam (review)"
    ),
    (
        "Week 6", "12–16 May", AMBER,
        "Phase 4: REST API Development",
        [
            "Build FastAPI service: POST /v1/score (single borrower — response <200ms)",
            "Build POST /v1/score/batch (bulk scoring — 1,000 borrowers in <10s)",
            "Build GET /v1/lenders/match and GET /v1/health endpoints",
            "Implement API key auth + Pydantic input validation + audit logging",
            "Containerise with Docker — model .pkl file baked into image",
            "Auto-generate OpenAPI/Swagger documentation for FinnUp Tech team",
        ],
        "Ganesh, Anil, Bhupesh, Hareram (API development leads)",
        "Remaining members: Asha, Arvind, Deepak, Gopal, Pranali, Rahul, Savitha, Sonam, Samik (testing + docs)"
    ),
    (
        "Week 7", "19–23 May", RED,
        "Phase 4: FinnUp Integration + UAT",
        [
            "Deploy Docker container to cloud (AWS/Azure) — live HTTPS URL provisioned",
            "Share API URL + credentials with FinnUp Tech team; support webhook configuration",
            "FinnUp maps scorecard fields (credit_score, risk_band, top_lenders) to dashboard",
            "UAT: run 200 known-outcome borrowers through live API; review scores with FinnUp team",
            "Validate SHAP explanations are interpretable to FinnUp loan officers",
            "Fix any issues found during UAT; obtain UAT sign-off from FinnUp",
        ],
        "Deepak, Asha, Sonam, Arvind (integration + UAT leads)",
        "Remaining members: Anil, Bhupesh, Ganesh, Gopal, Hareram, Pranali, Rahul, Savitha, Samik (support)"
    ),
    (
        "Week 8", "26–30 May", GREEN,
        "Phase 5: Final Report + IIM Calcutta Submission",
        [
            "Write final IIM Calcutta capstone report: problem, approach, results, business impact",
            "Compile all deliverables: model card, API documentation, UAT report, EDA report",
            "Prepare final presentation deck (15 slides) for IIM Calcutta review panel",
            "Document limitations, future work (monthly retraining, AA integration)",
            "Academic reflections: AI Canvas validation, ethical considerations, fairness analysis",
            "Submit all files to IIM Calcutta portal by May 31, 2026",
        ],
        "ALL 13 MEMBERS — Full team submission",
        "Asha, Arvind, Anil, Bhupesh, Deepak, Ganesh, Gopal, Hareram, Pranali, Rahul, Savitha, Sonam, Samik"
    ),
]

tbl_plan = doc.add_table(rows=1 + len(weeks_plan), cols=5)
tbl_plan.style = 'Table Grid'
for c, hdr in zip(tbl_plan.rows[0].cells,
                  ["Week", "Dates", "Phase / Theme", "Key Tasks (summary)", "Team Leads / All Members"]):
    set_cell_bg(c, NAVY)
    add_cell_para(c, hdr, bold=True, size=9.5, color=WHITE)

for i, (wk, dates, bg, theme, tasks, leads, rest) in enumerate(weeks_plan):
    row = tbl_plan.rows[i + 1]
    row_bg = LGRAY if i % 2 == 0 else WHITE

    set_cell_bg(row.cells[0], bg)
    add_cell_para(row.cells[0], wk, bold=True, size=10, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_bg(row.cells[1], row_bg)
    add_cell_para(row.cells[1], dates, size=9, color=GRAY)

    set_cell_bg(row.cells[2], row_bg)
    add_cell_para(row.cells[2], theme, bold=True, size=9.5, color=NAVY)

    set_cell_bg(row.cells[3], row_bg)
    tasks_str = "\n".join(f"  ▸ {t}" for t in tasks)
    row.cells[3].paragraphs[0].clear()
    r_t = row.cells[3].paragraphs[0].add_run(tasks_str)
    r_t.font.size = Pt(8.5); r_t.font.name = 'Calibri'
    r_t.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))

    set_cell_bg(row.cells[4], row_bg)
    row.cells[4].paragraphs[0].clear()
    r_l = row.cells[4].paragraphs[0].add_run(leads + "\n")
    r_l.bold = True; r_l.font.size = Pt(8.5); r_l.font.name = 'Calibri'
    r_l.font.color.rgb = RGBColor(*hex_to_rgb(TEAL))
    r_r = row.cells[4].paragraphs[0].add_run(rest)
    r_r.font.size = Pt(8); r_r.font.name = 'Calibri'; r_r.italic = True
    r_r.font.color.rgb = RGBColor(*hex_to_rgb(GRAY))

p()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — TECHNICAL ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
banner("PAGE 4 — TECHNICAL ARCHITECTURE: TWO-STAGE ML + API FLOW", bg=PURPLE)

# ── Two-stage model summary ────────────────────────────────────────────────────
mini_heading("Two-Stage ML Architecture")

tbl_arch = doc.add_table(rows=2, cols=2)
tbl_arch.style = 'Table Grid'

set_cell_bg(tbl_arch.rows[0].cells[0], TEAL)
add_cell_para(tbl_arch.rows[0].cells[0],
              "Stage 1 — Credit Scoring (XGBoost)", bold=True, size=10.5, color=WHITE)
set_cell_bg(tbl_arch.rows[0].cells[1], PURPLE)
add_cell_para(tbl_arch.rows[0].cells[1],
              "Stage 2 — Lender Ranking & Matching", bold=True, size=10.5, color=WHITE)

stage1_txt = (
    "Input: 50+ borrower features\n"
    "(Director CIBIL, Turnover, Bank Signals,\n"
    " Company Type, Geography, Loan Details)\n\n"
    "Model: XGBoost Classifier\n"
    "Training: 80% split, 5-fold CV\n"
    "Imbalance: SMOTE + scale_pos_weight\n"
    "Tuning: Optuna hyperparameter search\n\n"
    "Output: P(approval) score 0–1\n"
    "Calibrated to credit score 0–1000\n"
    "Risk bands: Poor / Fair / Good / Excellent\n\n"
    "Explainability: SHAP top-5 reasons per borrower"
)
stage2_txt = (
    "Input: Stage 1 score + lender policy rules\n\n"
    "Step 1: Rule engine\n"
    "  Filter lenders by hard eligibility rules:\n"
    "  CIBIL minimum, sector, state, loan size\n\n"
    "Step 2: Score each eligible lender\n"
    "  Run XGBoost for each lender separately\n"
    "  Get P(approval | this lender)\n\n"
    "Step 3: Rank and recommend\n"
    "  Sort eligible lenders high to low\n"
    "  Return top 3 as recommendations\n\n"
    "Output: Ranked lender list + score + reasons"
)

for col, txt in enumerate([stage1_txt, stage2_txt]):
    set_cell_bg(tbl_arch.rows[1].cells[col], LGRAY)
    tbl_arch.rows[1].cells[col].paragraphs[0].clear()
    r = tbl_arch.rows[1].cells[col].paragraphs[0].add_run(txt)
    r.font.size = Pt(9.5); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(*hex_to_rgb("1E293B"))

p()

# ── API end-to-end flow ────────────────────────────────────────────────────────
mini_heading("End-to-End API Flow: Borrower Registration → Lender Receives Scored Lead")

flow = [
    ("1", NAVY,   "Borrower registers on FinnUp portal"),
    ("2", TEAL,   "FinnUp sends raw data to POST /v1/score  (HTTPS + API key — <5ms)"),
    ("3", PURPLE, "Feature pipeline: 50+ features computed from raw JSON (~5ms)"),
    ("4", AMBER,  "Lender rule engine: filter to eligible lenders only (~1ms)"),
    ("5", TEAL,   "XGBoost Stage 1: P(approval) score per borrower (~3ms)"),
    ("6", PURPLE, "Stage 2 ranking: score each eligible lender, rank top 3 (~5ms)"),
    ("7", GREEN,  "SHAP: generate 'Top 5 reasons' explanation (~10ms)"),
    ("8", NAVY,   "API returns JSON scorecard: score, risk band, top lenders, reasons (~1ms)"),
    ("9", GREEN,  "FinnUp dashboard shows scored lead to matched lender"),
]

tbl_flow = doc.add_table(rows=len(flow), cols=2)
tbl_flow.style = 'Table Grid'
for i, (step, bg, desc) in enumerate(flow):
    set_cell_bg(tbl_flow.rows[i].cells[0], bg)
    add_cell_para(tbl_flow.rows[i].cells[0], f"Step {step}", bold=True, size=10,
                  color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    row_bg = LGRAY if i % 2 == 0 else WHITE
    set_cell_bg(tbl_flow.rows[i].cells[1], row_bg)
    add_cell_para(tbl_flow.rows[i].cells[1], desc, size=9.5, color="1E293B")

p()

# ── Timing summary bar ─────────────────────────────────────────────────────────
tbl_timing = doc.add_table(rows=1, cols=3)
tbl_timing.style = 'Table Grid'
for c, (label, val, bg) in enumerate(zip(
        ["Single borrower (real-time)", "1,000 borrowers (batch)", "Model loaded in memory"],
        ["< 80ms total response", "< 10 seconds total", "Once at API start — zero I/O per request"],
        [GREEN, TEAL, NAVY])):
    set_cell_bg(tbl_timing.rows[0].cells[c], bg)
    tbl_timing.rows[0].cells[c].paragraphs[0].clear()
    r1 = tbl_timing.rows[0].cells[c].paragraphs[0].add_run(val + "\n")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r1.font.color.rgb = RGBColor(*hex_to_rgb(WHITE))
    r2 = tbl_timing.rows[0].cells[c].paragraphs[0].add_run(label)
    r2.font.size = Pt(8.5); r2.font.name = 'Calibri'
    r2.font.color.rgb = RGBColor(*hex_to_rgb("DBEAFE"))
p()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — DELIVERABLES + OPEN ITEMS + TEAM
# ══════════════════════════════════════════════════════════════════════════════
banner("PAGE 5 — FINAL DELIVERABLES, OPEN ITEMS & TEAM", bg=NAVY)

# ── Final deliverables ─────────────────────────────────────────────────────────
mini_heading("Final Deliverables for IIM Calcutta (Due: May 31, 2026)")
deliverables = [
    ("◌", "EDA Report",               "EDA notebook (01_eda.ipynb) + auto-generated EDA PPT",                    GREEN,  "✓ DONE"),
    ("◌", "Feature Engineering",      "02_feature_engineering.ipynb — 50+ features, imputation pipeline",        AMBER,  "Week 2"),
    ("◌", "Trained ML Model",         "xgboost_model.pkl + feature_pipeline.pkl + lender_thresholds.pkl",        AMBER,  "Week 4"),
    ("◌", "Model Card",               "2-page formal doc: metrics, fairness analysis, limitations, use warnings", AMBER,  "Week 5"),
    ("◌", "SHAP Explanations",        "Top-5 reason generation per prediction — visual + API output",             AMBER,  "Week 5"),
    ("◌", "REST API",                 "FastAPI service — /v1/score, /v1/score/batch, /v1/lenders/match",          AMBER,  "Week 6"),
    ("◌", "API Documentation",        "Auto-generated OpenAPI/Swagger + integration guide for FinnUp Tech",       AMBER,  "Week 6"),
    ("◌", "Docker Container",         "Dockerfile + docker-compose — portable, cloud-deployable image",           AMBER,  "Week 6"),
    ("◌", "UAT Report",               "User Acceptance Test results — 200 borrowers, FinnUp team sign-off",       AMBER,  "Week 7"),
    ("◌", "Project Roadmap PPT",      "FinnUp_ProjectStatus_Roadmap.pptx — 15-slide stakeholder deck",           AMBER,  "Week 8"),
    ("◌", "IIM Submission PPT",       "FinnUp_IIM_Submission.pptx — 15-slide academic submission deck",          AMBER,  "Week 8"),
    ("◌", "Final Capstone Report",    "Full written report: problem, approach, results, business impact, ethics", AMBER,  "Week 8"),
    ("◌", "Academic Reflection",      "AI Canvas validation, ethical considerations, fairness analysis section",  AMBER,  "Week 8"),
]

tbl_del = doc.add_table(rows=1 + len(deliverables), cols=4)
tbl_del.style = 'Table Grid'
for c, hdr in zip(tbl_del.rows[0].cells, ["Deliverable", "Description", "Due", "Status"]):
    set_cell_bg(c, NAVY)
    add_cell_para(c, hdr, bold=True, size=9.5, color=WHITE)

for i, (_, name, desc, col, status) in enumerate(deliverables):
    row = tbl_del.rows[i + 1]
    bg_d = LGREEN if col == GREEN else (LAMBER if col == AMBER else LGRAY)
    set_cell_bg(row.cells[0], bg_d)
    add_cell_para(row.cells[0], name, bold=True, size=9.5, color=NAVY)
    set_cell_bg(row.cells[1], bg_d)
    add_cell_para(row.cells[1], desc, size=9, color="1E293B")
    set_cell_bg(row.cells[2], bg_d)
    add_cell_para(row.cells[2], status.replace("✓ DONE", "").strip() if "Week" in status else "",
                  size=9, color=GRAY)
    set_cell_bg(row.cells[3], bg_d)
    status_color = GREEN if "DONE" in status else AMBER
    add_cell_para(row.cells[3], status, bold=True, size=9, color=status_color)

p()

# ── Open blockers ──────────────────────────────────────────────────────────────
mini_heading("Open Blockers (Action Required — Week 1)")
blockers = [
    ("CRITICAL", RED,
     "Outcome labels not yet collected — FinnUp must provide approved/rejected history "
     "for as many of the 6,483 borrowers as possible. Without this, FinnUp Status proxy "
     "will be used and model quality may be lower."),
    ("CRITICAL", RED,
     "Bank statement coverage only 38% (2,489 of 6,483 borrowers) — "
     "Account Aggregator (AA) consent process must be initiated for uncovered borrowers. "
     "Even partial improvement directly improves model accuracy."),
    ("HIGH", AMBER,
     "Lender policy files must be confirmed as complete and current — any outdated eligibility "
     "rules will cause incorrect lender filtering in Stage 2."),
    ("MEDIUM", AMBER,
     "Cloud environment (AWS/Azure) must be provisioned by Week 5 to allow dockerised API "
     "deployment in Week 6 and FinnUp integration in Week 7."),
]

tbl_blk = doc.add_table(rows=1 + len(blockers), cols=3)
tbl_blk.style = 'Table Grid'
for c, hdr in zip(tbl_blk.rows[0].cells, ["Priority", "Blocker", "Owner"]):
    set_cell_bg(c, RED)
    add_cell_para(c, hdr, bold=True, size=9.5, color=WHITE)

owners_blk = ["Ganesh + FinnUp Team", "Hareram + FinnUp Team",
              "Bhupesh + Savitha", "Deepak + Arvind"]
for i, ((priority, col, desc), owner) in enumerate(zip(blockers, owners_blk)):
    row = tbl_blk.rows[i + 1]
    bg_b = LRED if col == RED else LAMBER
    set_cell_bg(row.cells[0], bg_b)
    add_cell_para(row.cells[0], priority, bold=True, size=9.5, color=col)
    set_cell_bg(row.cells[1], bg_b)
    add_cell_para(row.cells[1], desc, size=9, color="1E293B")
    set_cell_bg(row.cells[2], bg_b)
    add_cell_para(row.cells[2], owner, size=9, color=NAVY)

p()
divider()
p()

# ── Full team table ────────────────────────────────────────────────────────────
mini_heading("Full Team — Cohort 2, Group 1 (13 Members)")

members = [
    ("Asha",    "EDA review, feature validation, UAT testing"),
    ("Arvind",  "Lender policy review, integration support, final report"),
    ("Anil",    "Feature engineering, API development, code review"),
    ("Bhupesh", "Data augmentation, lender policy confirmation, API testing"),
    ("Deepak",  "Label creation, feature selection, FinnUp integration lead"),
    ("Ganesh",  "Data lead, model training, project coordination"),
    ("Gopal",   "Model training, Stage 2 lender ranking, results review"),
    ("Hareram", "Bank data / AA consent, API development, infrastructure"),
    ("Pranali", "Feature engineering lead, imputation pipeline, peer review"),
    ("Rahul",   "Model training lead, XGBoost + SHAP, model card author"),
    ("Savitha", "Data validation, model card review, final report"),
    ("Sonam",   "Label creation, feature engineering, FinnUp integration support"),
    ("Samik",   "Stage 2 ranking, end-to-end testing, presentation prep"),
]

cols_m = 2
rows_m = (len(members) + cols_m - 1) // cols_m
tbl_team2 = doc.add_table(rows=rows_m + 1, cols=cols_m * 2)
tbl_team2.style = 'Table Grid'
for c in range(cols_m):
    set_cell_bg(tbl_team2.rows[0].cells[c * 2], TEAL)
    add_cell_para(tbl_team2.rows[0].cells[c * 2], "Member", bold=True, size=9.5, color=WHITE)
    set_cell_bg(tbl_team2.rows[0].cells[c * 2 + 1], TEAL)
    add_cell_para(tbl_team2.rows[0].cells[c * 2 + 1], "Primary Focus Areas", bold=True, size=9.5, color=WHITE)

for idx, (name, role) in enumerate(members):
    row_i = (idx // cols_m) + 1
    col_i = (idx % cols_m) * 2
    bg_m = LGRAY if row_i % 2 == 0 else WHITE
    set_cell_bg(tbl_team2.rows[row_i].cells[col_i], bg_m)
    add_cell_para(tbl_team2.rows[row_i].cells[col_i], name, bold=True, size=9.5, color=NAVY)
    set_cell_bg(tbl_team2.rows[row_i].cells[col_i + 1], bg_m)
    add_cell_para(tbl_team2.rows[row_i].cells[col_i + 1], role, size=9, color="1E293B")

p()

# ── Footer strip ───────────────────────────────────────────────────────────────
tbl_footer = doc.add_table(rows=1, cols=1)
tc_f = tbl_footer.rows[0].cells[0]
set_cell_bg(tc_f, NAVY)
tc_f.paragraphs[0].clear()
tc_f.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = tc_f.paragraphs[0].add_run(
    "FinnUp MSME Credit Intelligence  |  APAL Cohort 2 – Group 1  |  IIM Calcutta  |  "
    "Target Completion: May 31, 2026  |  Confidential"
)
rf.font.size = Pt(8.5); rf.font.name = 'Calibri'
rf.font.color.rgb = RGBColor(*hex_to_rgb("DBEAFE"))

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"✓  Saved: {OUTPUT_PATH.resolve()}")
